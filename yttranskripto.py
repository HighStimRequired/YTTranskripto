import sys
import csv
import json
from datetime import datetime

from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import QSettings
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLineEdit, QPushButton, QTextEdit, QLabel, QFileDialog, QComboBox, QCheckBox, QMessageBox
)
# For DOCX export; install via: pip install python-docx
from docx import Document

# External libraries: install via pip if needed:
# pip install youtube-transcript-api pytube PyQt5 python-docx
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
from pytube import YouTube

# ---------------- Utility Functions ----------------

def extract_video_id(url):
    """Extract the YouTube video ID from a URL."""
    if "youtube.com" in url:
        parts = url.split("v=")
        if len(parts) > 1:
            return parts[1].split("&")[0]
    elif "youtu.be" in url:
        parts = url.split("/")
        if parts:
            return parts[-1]
    return None

def safe_time(value):
    """
    Recursively extract a numeric value from a nested structure.
    If 'value' is a dict and contains a key "value", return its numeric conversion.
    Otherwise, iterate through nested values.
    """
    if isinstance(value, dict):
        if "value" in value:
            try:
                return float(value["value"])
            except (ValueError, TypeError):
                pass
        for v in value.values():
            num = safe_time(v)
            if isinstance(num, (int, float)):
                return num
        return 0.0
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0

def format_timestamp(seconds, fmt="HH:MM:SS"):
    """Format seconds into a timestamp string based on selected format."""
    seconds = int(safe_time(seconds))
    if fmt == "HH:MM:SS":
        return datetime.utcfromtimestamp(seconds).strftime("%H:%M:%S")
    elif fmt == "mm:ss":
        minutes = seconds // 60
        secs = seconds % 60
        return f"{minutes:02}:{secs:02}"
    else:
        return datetime.utcfromtimestamp(seconds).strftime("%H:%M:%S")

def format_srt_timestamp(seconds):
    """Format seconds into SRT timestamp (HH:MM:SS,mmm)."""
    seconds = safe_time(seconds)
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    msecs = int(round((seconds - int(seconds)) * 1000))
    return f"{hours:02}:{minutes:02}:{secs:02},{msecs:03}"

# ---------------- Main Application ----------------

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("YTTranskripto")
        self.setGeometry(100, 100, 900, 600)
        self.transcript_data = None  # List to store transcript entries

        # Setup QSettings to save user preferences (for export format, timestamps, and timestamp format)
        self.settings = QSettings("MyCompany", "YTTranskripto")
        self.default_export_format = self.settings.value("default_export_format", "TXT")
        self.include_timestamps = self.settings.value("include_timestamps", "true") == "true"
        self.timestamp_format = self.settings.value("timestamp_format", "HH:MM:SS")

        self.init_ui()

    def init_ui(self):
        main_widget = QWidget()
        main_layout = QVBoxLayout()

        # ---- URL Input and Transcript Retrieval ----
        url_layout = QHBoxLayout()
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Enter YouTube URL here")
        self.get_transcript_btn = QPushButton("Get Transcript")
        self.get_transcript_btn.clicked.connect(self.fetch_transcript)
        url_layout.addWidget(self.url_input)
        url_layout.addWidget(self.get_transcript_btn)
        main_layout.addLayout(url_layout)

        # ---- Preferences: Timestamp Options ----
        ts_layout = QHBoxLayout()
        self.timestamp_checkbox = QCheckBox("Include Timestamps")
        self.timestamp_checkbox.setChecked(self.include_timestamps)
        # Update display on state change.
        self.timestamp_checkbox.stateChanged.connect(self.update_transcript_display)
        ts_layout.addWidget(self.timestamp_checkbox)

        # Add a spacer to push the label to the right
        ts_layout.addStretch(1)

        # Create a right-aligned label for "Timestamp Format:" with increased fixed width
        label_format = QLabel("Timestamp Format:")
        label_format.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter)
        label_format.setFixedWidth(150)
        ts_layout.addWidget(label_format)

        self.timestamp_format_combo = QComboBox()
        self.timestamp_format_combo.addItems(["HH:MM:SS", "mm:ss"])
        index = self.timestamp_format_combo.findText(self.timestamp_format)
        if index >= 0:
            self.timestamp_format_combo.setCurrentIndex(index)
        # Update display on format change.
        self.timestamp_format_combo.currentIndexChanged.connect(self.update_transcript_display)
        ts_layout.addWidget(self.timestamp_format_combo)
        main_layout.addLayout(ts_layout)

        # ---- Transcript Display ----
        self.transcript_display = QTextEdit()
        self.transcript_display.setReadOnly(True)
        main_layout.addWidget(self.transcript_display)

        # ---- Export Options ----
        export_layout = QHBoxLayout()
        self.export_format_combo = QComboBox()
        self.export_format_combo.addItems(["TXT", "CSV", "JSON", "DOCX", "SRT"])
        index = self.export_format_combo.findText(self.default_export_format)
        if index >= 0:
            self.export_format_combo.setCurrentIndex(index)
        export_layout.addWidget(QLabel("Export as:"))
        export_layout.addWidget(self.export_format_combo)
        self.export_btn = QPushButton("Export Transcript")
        self.export_btn.clicked.connect(self.export_transcript)
        export_layout.addWidget(self.export_btn)
        main_layout.addLayout(export_layout)

        main_widget.setLayout(main_layout)
        self.setCentralWidget(main_widget)
        self.apply_dark_theme()

    def fetch_transcript(self):
        """Fetch the transcript for the given YouTube URL.
           The app automatically tries to fetch the English transcript ("en"),
           falling back to the first available transcript if necessary.
        """
        self.transcript_display.clear()
        url = self.url_input.text().strip()
        if not url:
            QMessageBox.warning(self, "Input Error", "Please enter a valid YouTube URL.")
            return

        video_id = extract_video_id(url)
        if not video_id:
            QMessageBox.warning(self, "URL Error", "Could not extract video ID from URL.")
            return

        try:
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            try:
                transcript_obj = transcript_list.find_transcript(["en"])
            except Exception:
                # Fallback to the first available transcript
                transcript_obj = transcript_list.find_transcript([transcript_list._generated_transcripts[0].language_code])
            self.transcript_data = transcript_obj.fetch()
            self.update_transcript_display()
        except (TranscriptsDisabled, NoTranscriptFound):
            QMessageBox.information(self, "Transcript Unavailable", "Transcript not available for this video.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error fetching transcript:\n{str(e)}")

    def update_transcript_display(self):
        """Re-render the transcript display based on the current timestamp settings."""
        if not self.transcript_data:
            return
        include_ts = self.timestamp_checkbox.isChecked()
        ts_format = self.timestamp_format_combo.currentText()
        display_text = ""
        for entry in self.transcript_data:
            if include_ts:
                ts = format_timestamp(entry["start"], ts_format)
                display_text += f"[{ts}] {entry['text']}\n"
            else:
                display_text += f"{entry['text']}\n"
        self.transcript_display.setText(display_text)

    def export_transcript(self):
        """Export the fetched transcript in the selected format, using the timestamp options."""
        if not self.transcript_data:
            QMessageBox.warning(self, "No Transcript", "No transcript to export. Please fetch a transcript first.")
            return

        export_format = self.export_format_combo.currentText().lower()
        options = QFileDialog.Options()
        file_filter = ""
        if export_format == "txt":
            file_filter = "Text Files (*.txt)"
        elif export_format == "csv":
            file_filter = "CSV Files (*.csv)"
        elif export_format == "json":
            file_filter = "JSON Files (*.json)"
        elif export_format == "docx":
            file_filter = "Word Documents (*.docx)"
        elif export_format == "srt":
            file_filter = "SubRip Files (*.srt)"

        file_path, _ = QFileDialog.getSaveFileName(self, "Save Transcript", "", file_filter, options=options)
        if not file_path:
            return

        include_ts = self.timestamp_checkbox.isChecked()
        ts_format = self.timestamp_format_combo.currentText()

        try:
            if export_format == "txt":
                with open(file_path, "w", encoding="utf-8") as f:
                    for entry in self.transcript_data:
                        if include_ts:
                            ts = format_timestamp(entry["start"], ts_format)
                            f.write(f"[{ts}] {entry['text']}\n")
                        else:
                            f.write(f"{entry['text']}\n")
            elif export_format == "csv":
                with open(file_path, "w", newline="", encoding="utf-8") as csvfile:
                    writer = csv.writer(csvfile)
                    writer.writerow(["Start Time", "Text", "Duration"])
                    for entry in self.transcript_data:
                        if include_ts:
                            ts = format_timestamp(entry["start"], ts_format)
                        else:
                            ts = ""
                        writer.writerow([ts, entry["text"], entry["duration"]])
            elif export_format == "json":
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(self.transcript_data, f, indent=4)
            elif export_format == "docx":
                document = Document()
                document.add_heading("Transcript", level=1)
                for entry in self.transcript_data:
                    if include_ts:
                        ts = format_timestamp(entry["start"], ts_format)
                        paragraph_text = f"[{ts}] {entry['text']}"
                    else:
                        paragraph_text = entry["text"]
                    document.add_paragraph(paragraph_text)
                document.save(file_path)
            elif export_format == "srt":
                with open(file_path, "w", encoding="utf-8") as f:
                    for i, entry in enumerate(self.transcript_data, start=1):
                        start_val = safe_time(entry["start"])
                        duration_val = safe_time(entry["duration"])
                        start_srt = format_srt_timestamp(start_val)
                        end_time = start_val + duration_val
                        end_srt = format_srt_timestamp(end_time)
                        f.write(f"{i}\n")
                        f.write(f"{start_srt} --> {end_srt}\n")
                        f.write(f"{entry['text']}\n\n")
            QMessageBox.information(self, "Export Successful", f"Transcript exported successfully to {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Error exporting transcript:\n{str(e)}")

        self.save_preferences()

    def save_preferences(self):
        """Save current user preferences using QSettings."""
        self.settings.setValue("default_export_format", self.export_format_combo.currentText())
        self.settings.setValue("include_timestamps", "true" if self.timestamp_checkbox.isChecked() else "false")
        self.settings.setValue("timestamp_format", self.timestamp_format_combo.currentText())

    def apply_dark_theme(self):
        """Apply a cozy dark theme with YouTube-red accents and softened UI elements."""
        dark_stylesheet = """
            QWidget {
                background-color: #121212;
                color: #e0e0e0;
                font-family: 'Helvetica', Arial, sans-serif;
                font-size: 14px;
                padding: 5px;
            }
            QLineEdit, QTextEdit, QComboBox, QCheckBox {
                background-color: #1e1e1e;
                border: 1px solid #333;
                color: #e0e0e0;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton {
                background-color: #FF0000;
                border: none;
                padding: 5px 10px;
                color: #fff;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #cc0000;
            }
        """
        self.setStyleSheet(dark_stylesheet)

    def closeEvent(self, event):
        """Save preferences before closing."""
        self.save_preferences()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
